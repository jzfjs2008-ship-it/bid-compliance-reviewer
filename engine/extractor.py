import re
import fitz
from docx import Document
from openpyxl import load_workbook


class Extractor:
    @staticmethod
    def extract_pdf(file_path):
        pages = []
        doc = fitz.open(file_path)
        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                blocks = page.get_text("dict")["blocks"]
                rich_spans = []
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                span_text = span.get("text", "").strip()
                                if not span_text:
                                    continue
                                is_bold = (span.get("font", "") and "bold" in span.get("font", "").lower())
                                color = span.get("color", 0)
                                if is_bold or color != 0:
                                    span_text = f"[BOLD]{span_text}"
                                rich_spans.append(span_text)
                combined_rich = " ".join(rich_spans) if rich_spans else text
                pages.append({
                    "page": page_num + 1,
                    "text": text,
                    "rich_text": combined_rich,
                })
        finally:
            doc.close()
        return pages

    @staticmethod
    def extract_docx(file_path):
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                "style": para.style.name if para.style else "Normal",
                "text": para.text,
                "runs": [{"text": r.text, "bold": r.bold} for r in para.runs],
            })
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        return {"paragraphs": paragraphs, "tables": tables}

    @staticmethod
    def extract_excel(file_path):
        wb = load_workbook(file_path, read_only=True, data_only=True)
        sheets = {}
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(c) if c is not None else "" for c in row])
            sheets[sheet_name] = rows
        wb.close()
        return sheets

    @staticmethod
    def extract_pdf_stream(file_path, rules, queue_handler=None):
        results = {}
        doc = fitz.open(file_path)
        try:
            for rule in rules:
                field_name = rule["field_name"]
                anchor = rule["anchor"]
                regex = rule["regex"]
                pattern = re.compile(regex)
                found = None
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text("text")
                    anchor_pos = text.find(anchor)
                    if anchor_pos == -1:
                        continue
                    match = pattern.search(text)
                    if match:
                        found = match.group(1).strip()
                        if queue_handler:
                            queue_handler.put("progress", f"提取字段 [{field_name}]: {found}")
                        break
                results[field_name] = found
        finally:
            doc.close()
        return results

    @staticmethod
    def extract_table_from_pdf(file_path):
        doc = fitz.open(file_path)
        tables = []
        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                tabs = page.find_tables()
                for tab in tabs:
                    rows = []
                    for row in tab.extract():
                        rows.append([str(c) if c else "" for c in row])
                    if rows:
                        tables.append({"page": page_num + 1, "data": rows})
        finally:
            doc.close()
        return tables

    @staticmethod
    def extract_docx_table_data(file_path):
        doc = Document(file_path)
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)
        return tables
